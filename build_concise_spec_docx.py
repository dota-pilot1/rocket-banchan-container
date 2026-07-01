# -*- coding: utf-8 -*-
"""Create an editable DOCX version of the concise EnglishAgentHub spec."""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "/Users/terecal/english-agent-hub-container/output/docx/EnglishAgentHub_스펙명세서_간결판_편집가능.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "334155"
MUTED = "64748B"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "ECFDF5"
CALLOUT_BAR = "059669"
BORDER = "CBD5E1"


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=11, color=TEXT, bold=False):
    for run in paragraph.runs:
        set_font(run, size=size, color=color, bold=bold)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(int(w * 1440) for w in widths_in)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = int(widths_in[idx] * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.allow_autofit = False
    header = table.rows[0].cells
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, text in enumerate(headers):
        header[i].text = text
        set_cell_shading(header[i], TABLE_FILL)
        for p in header[i].paragraphs:
            set_paragraph_font(p, size=9.5, color=DARK_BLUE, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            if r_idx % 2 == 0:
                set_cell_shading(cells[i], LIGHT_FILL)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                set_paragraph_font(p, size=9.2, color=TEXT)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_kv(doc, rows, label_width=1.875):
    return add_table(doc, ["항목", "내용"], rows, [label_width, 6.5 - label_width])


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    set_paragraph_font(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    set_paragraph_font(p)
    return p


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_shading(cell, CALLOUT_FILL)
    for p in cell.paragraphs:
        set_paragraph_font(p, size=9.5, color=TEXT)
        p.paragraph_format.space_after = Pt(0)
    set_table_geometry(table, [6.5])
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.start_type = WD_SECTION.NEW_PAGE
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.line_spacing = 1.25

for s_name in ("List Bullet", "List Number"):
    s = styles[s_name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    s.font.size = Pt(11)
    s.paragraph_format.left_indent = Inches(0.375)
    s.paragraph_format.first_line_indent = Inches(-0.188)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.25

footer_p = section.footer.paragraphs[0]
footer_p.text = "EnglishAgentHub · 기술 스펙 간결판"
set_paragraph_font(footer_p, size=8, color=MUTED)
footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("EnglishAgentHub")
set_font(r, size=26, bold=True, color=BLUE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("기술 스펙 명세서 - 간결판")
set_font(r, size=16, bold=False, color=TEXT)
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tagline.add_run("AI 영어 학습 · 문제 은행 · 시험 출제/응시 플랫폼")
set_font(r, size=11, color=MUTED)

add_kv(
    doc,
    [
        ("도메인", "dxline-tallent.com"),
        ("Frontend", "Next.js 16 · React 19 · TypeScript · Tailwind v4"),
        ("Backend", "Spring Boot 4 · Java 21 · Spring AI"),
        ("DB", "PostgreSQL + pgvector"),
        ("인증", "JWT Stateless · RBAC · 사용자별 OpenAI API Key 암호화"),
        ("작성일", "2026-06-23"),
    ],
    label_width=1.181,
)

add_heading(doc, "1. 개요")
add_body(
    doc,
    "EnglishAgentHub는 AI 영어 학습과 문제 은행/시험 운영을 한곳에서 제공하는 웹 플랫폼이다. "
    "영어 회화 에이전트, 번역, 표현 피드백, 음성 학습을 제공하고, PDF/이미지 문항 추출부터 시험지 생성, 응시, 자동 채점, 결과 분석, 오답노트까지 지원한다.",
)
for item in [
    "AI 학습: 에이전트 채팅, 한영/영한 번역, 표현 피드백, 답변 추천, 뉴스 학습, TTS/STT, Realtime Voice",
    "문제 은행: 문항 CRUD, AI 유사 문항 생성, PDF/수학 문항 추출, pgvector 유사 검색",
    "시험: 시험지 생성/게시/마감, 응시 재개, 자동 채점, 결과 분석, 오답노트",
    "관리: 회원/프로필, 사용자 OpenAI API Key, 역할/권한, 메뉴/사이트 설정",
]:
    add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "2. 기술 스택")
add_kv(
    doc,
    [
        ("프론트엔드", "Next.js 16.1.4, React 19.2.3, TypeScript 5, Tailwind CSS v4, TanStack Query, Axios, Zod, react-hook-form"),
        ("UI/기능", "Radix UI, AG Grid, Framer Motion, dnd-kit, KaTeX, i18next, socket.io-client"),
        ("백엔드", "Spring Boot 4.0.5, Java 21, Gradle, Spring Security, JPA/Hibernate, SpringDoc OpenAPI"),
        ("AI/문서/저장", "Spring AI OpenAI 2.0.0-M7, PDFBox 3.0.3, AWS S3 SDK 2.29"),
        ("데이터베이스", "PostgreSQL(localhost:5436/english_agent_hub), pgvector, 로컬 계정 postgres/postgres"),
    ],
)

add_heading(doc, "3. 아키텍처")
add_table(
    doc,
    ["구성", "역할"],
    [
        ("Frontend", "Next.js 정적 빌드 · S3/CloudFront · dev 3300"),
        ("Reverse Proxy", "Nginx가 /api 요청을 Spring Boot로 전달"),
        ("Backend", "Spring Boot API · local 3301 / prod 4301 · Spring AI 통해 OpenAI 호출"),
        ("Database", "PostgreSQL 5436 + pgvector · 문항 임베딩 및 유사 검색"),
    ],
    [1.875, 4.625],
)

add_heading(doc, "4. 인증/보안")
add_kv(
    doc,
    [
        ("인증 방식", "Spring Security + JWT Stateless. Authorization: Bearer 토큰을 요청마다 검증"),
        ("토큰", "Access 30분, Refresh 7일. Refresh는 /api/auth/refresh에서 회전"),
        ("권한", "Role/Permission N:M 구조, @EnableMethodSecurity 및 @PreAuthorize 적용"),
        ("공개 API", "/api/auth/signup, /login, /refresh, /check-email, GET /api/agents/**, /api/site-settings, /api/menus, Swagger/OpenAPI, OPTIONS"),
        ("API Key", "사용자 OpenAI API Key는 AES-GCM 256bit로 암호화 저장. /api/users/me/api-key에서 조회(마스킹)/저장/삭제/검증"),
    ],
)
add_table(
    doc,
    ["Base", "주요 엔드포인트"],
    [
        ("/api/auth", "POST /signup, /login, /refresh, /logout · GET /check-email, /me"),
        ("/api/users/me/api-key", "GET 마스킹 조회 · PUT 저장 · DELETE 삭제 · POST /openai/validate"),
    ],
    [1.875, 4.625],
)

doc.add_page_break()
add_heading(doc, "5. 핵심 도메인 API")
add_callout(doc, "관리성 API는 기본적으로 ADMIN 권한으로 보호된다. 학생 응시/결과 조회는 인증 사용자 기준으로 제한한다.")
add_table(
    doc,
    ["도메인", "Base API", "핵심 기능"],
    [
        ("문항", "/api/questions", "목록/상세/생성/수정/삭제, 임베딩 생성, 유사 문항 검색, AI 유사 독해 문항 생성"),
        ("카테고리", "/api/categories", "self-reference 트리 기반 문항 분류"),
        ("PDF 독해 추출", "/api/extracted-sheets", "PDF 업로드, 독해 문항 추출, 추출 시트 조회/삭제"),
        ("수학 이미지 추출", "/api/extracted-math-sheets", "수학 PDF 문항 이미지 분할, answerFile 옵션, 시트 조회/삭제"),
        ("정형 수학 추출", "/api/structured-math-sheets", "202 + jobId 비동기 추출, /jobs/{jobId} 폴링, LaTeX 전사 및 도형 분리"),
        ("시험지", "/api/exams", "DRAFT 생성/수정, AI 변형본 생성, 게시(PUBLISHED), 마감(CLOSED), 삭제"),
        ("응시/채점", "/api/attempts", "응시 시작/재개, 제출 및 자동 채점, 결과/이력 조회, 관리자 응시 내역 관리"),
        ("공개 시험", "/api/practice/exams", "PUBLISHED 시험만 학생에게 노출"),
    ],
    [1.15, 1.75, 3.6],
)
add_heading(doc, "시험 흐름", level=2)
add_body(doc, "문항 등록/추출 → 임베딩/유사 검색 → 시험지 구성(DRAFT) → 게시(PUBLISHED) → 응시 시작/제출 → 자동 채점 → 결과 분석/오답노트 → 마감(CLOSED)")

add_heading(doc, "6. 기타 API")
add_table(
    doc,
    ["영역", "Controller / API"],
    [
        ("AI 채팅/학습", "AiChatController, LearningAgentController, RealtimeController (/api/ai, /api/agents, /api/realtime)"),
        ("어휘", "EnglishVocabularyController"),
        ("파일", "UploadController: S3 업로드 및 presigned URL"),
        ("사용자/RBAC", "UserManagementController, UserApiKeyController, RoleController, PermissionController, PermissionCategoryController"),
        ("UI/캐릭터", "MenuController (/api/menus), SiteSettingController (/api/site-settings), CharacterController: AI 에이전트 NPC"),
    ],
    [1.45, 5.05],
)

add_heading(doc, "7. 배포/외부 연동")
add_kv(
    doc,
    [
        ("운영", "Frontend S3 + CloudFront(443), Backend EC2 + systemd(4301), Nginx Reverse Proxy, PostgreSQL Docker + pgvector(5436)"),
        ("AWS S3", "region ap-northeast-2, prefix english-agent-hub, presign 300s"),
        ("OpenAI", "Chat gpt-5-mini, Embedding text-embedding-3-small, 번역 gpt-5-nano, STT gpt-4o-mini-transcribe, Realtime gpt-realtime-2 voice marin"),
        ("업로드 제한", "multipart 25MB"),
        ("주의", "BeautyBook에서 인계된 인프라라 일부 S3 버킷명/잔재에 beauty-book 흔적이 있어도 정상"),
    ],
)

doc.core_properties.title = "EnglishAgentHub 기술 스펙 명세서 간결판"
doc.core_properties.author = "EnglishAgentHub"
doc.save(OUT)
print(OUT)
