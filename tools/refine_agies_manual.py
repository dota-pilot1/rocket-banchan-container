from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE = Path("/Users/terecal/Desktop/agies_api_server_menual.docx")
OUTPUT = Path("/Users/terecal/Desktop/agies_api_server_manual_refined.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F6F8FB"
BORDER = "CAD3DF"
CODE_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, col_widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=None, bold=None, color=None, font="Malgun Gothic") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_rule(paragraph, color=BLUE, size="12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_style_font(style, font="Malgun Gothic", size=11, color=INK, bold=False) -> None:
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    set_style_font(styles["Normal"], size=10.5, color=INK)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.22

    set_style_font(styles["Heading 1"], size=16, color=BLUE, bold=True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], size=13, color=BLUE, bold=True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], size=11.5, color=DARK_BLUE, bold=True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        set_style_font(style, size=10.5, color=INK)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.22

    code_style = styles.add_style("EAH Code", 1)
    set_style_font(code_style, font="Consolas", size=8.8, color="263238")
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.05

    note_style = styles.add_style("EAH Note", 1)
    set_style_font(note_style, size=10.2, color=INK)
    note_style.paragraph_format.space_after = Pt(4)
    note_style.paragraph_format.line_spacing = 1.18

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("EnglishAgentHub API Server Manual")
    style_run(r, size=8.5, bold=True, color=MUTED)
    add_rule(header, color=BORDER, size="4")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("EnglishAgentHub · 2026-06-23")
    style_run(r, size=8.5, color=MUTED)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run("EnglishAgentHub")
    style_run(r, size=26, bold=True, color="0B2545")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.1)
    r = p.add_run("API 서버 기술 스펙 명세서")
    style_run(r, size=15, bold=True, color=BLUE)
    add_rule(p, color=BLUE, size="16")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=10, line=1.2)
    r = p.add_run("AI 영어 학습 · 문제 은행 · 시험 출제/응시 플랫폼")
    style_run(r, size=10.8, color=MUTED)


def add_table(doc: Document, rows: list[list[str]], widths: list[float], header=True) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        if r_idx == 0 and header:
            set_repeat_table_header(table.rows[0])
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.12)
            for part_idx, part in enumerate(text.split("\n")):
                if part_idx:
                    p.add_run().add_break()
                run = p.add_run(part)
                style_run(run, size=9.3 if len(row) >= 3 else 9.7, bold=(header and r_idx == 0), color=INK)
            set_cell_margins(cell)
            set_cell_borders(cell)
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif c_idx == 0 and len(row) == 2:
                set_cell_shading(cell, LIGHT_GRAY)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    doc.add_paragraph()


def add_note_box(doc: Document, title: str, body: str, fill=LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="D6DEE8")
    set_cell_margins(cell, top=110, bottom=110, start=160, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.15)
    r = p.add_run(title)
    style_run(r, size=10.2, bold=True, color=DARK_BLUE)
    if body:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, after=0, line=1.18)
        r = p.add_run(body)
        style_run(r, size=9.7, color=INK)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_FILL)
    set_cell_borders(cell, color="D8DEE8")
    set_cell_margins(cell, top=100, bottom=100, start=150, end=150)
    cell.text = ""
    for idx, line in enumerate(code.strip("\n").splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "EAH Code"
        r = p.add_run(line)
        style_run(r, font="Consolas", size=8.8, color="263238")
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        style_run(r, size=10.4, color=INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        add_rule(p, color="D9E4F1", size="6")


def add_api_flow(doc: Document) -> None:
    add_heading(doc, "응답 예시", 2)
    steps = [
        ("1", "잡 시작", "POST /api/structured-math-sheets", "202 Accepted + jobId"),
        ("2", "상태 폴링", "GET /api/structured-math-sheets/jobs/{jobId}", "PROCESSING · DONE · FAILED"),
        ("3", "시험지 단건 조회", "GET /api/structured-math-sheets/{id}", "문항 배열 포함"),
    ]
    add_table(doc, [["순서", "동작", "요청", "응답"]] + [list(s) for s in steps], [0.55, 1.25, 3.1, 1.6])

    add_code_block(doc, '{ "jobId": "f47ac10b-58cc-4372-a567-0e02b2c3d479" }')
    add_code_block(doc, '{ "status": "PROCESSING", "sheetId": null, "error": null, "total": 46, "done": 12 }')
    add_note_box(doc, "완료 상태", 'status가 "DONE"이면 sheetId가 채워지고, "FAILED"이면 error에 실패 사유가 들어갑니다.', fill="F7FAFF")

    add_heading(doc, "시험지 단건 응답", 3)
    add_code_block(
        doc,
        """{
  "id": "b5241cb8-0bec-4c47-8047-28933f28f918",
  "title": "2025학년도-대학수학능력시험-수학-문제",
  "sourceFileName": "2025_suneung_math.pdf",
  "itemCount": 46,
  "createdByName": "오현석",
  "createdAt": "2026-06-23T05:12:30Z",
  "items": [{
    "questionNumber": 1,
    "prompt": "$\\\\sqrt[3]{5} \\\\times 25^{\\\\frac{1}{3}}$ 의 값은?",
    "choices": ["1", "2", "3", "4", "5"],
    "figureImageUrl": null,
    "answer": "5",
    "points": 2,
    "type": "객관식",
    "subject": "수학",
    "needsReview": false
  }]
}""",
    )


def build() -> None:
    source_doc = Document(SOURCE)
    source_tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in source_doc.tables
    ]

    doc = Document()
    setup_document(doc)
    add_title_block(doc)
    add_table(doc, source_tables[0], [1.35, 5.15])

    add_heading(doc, "1. 개요", 1)
    add_bullets(
        doc,
        [
            "문제 은행, 시험 생성, 조회 API를 중심으로 구성된 영어 학습 플랫폼입니다.",
            "문제 은행: 문항 CRUD, AI 유사 문항 생성, PDF/수학 문항 추출, pgvector 유사 검색",
            "시험: 시험지 생성/게시/마감, 응시 재개, 자동 채점, 결과 분석, 오답노트",
            "관리: 회원/프로필, 사용자 OpenAI API Key, 역할/권한, 메뉴/사이트 설정",
        ],
    )

    add_heading(doc, "2. 기술 스택", 1)
    add_table(doc, source_tables[1], [1.45, 5.05])

    add_heading(doc, "3. 아키텍처", 1)
    add_table(doc, source_tables[2], [1.45, 5.05])

    add_heading(doc, "4. 인증/보안", 1)
    add_table(doc, source_tables[3], [1.45, 5.05])
    add_note_box(doc, "운영 URL", "http://dxline-tallent.com/", fill="F7FAFF")
    add_heading(doc, "인증 관련 엔드포인트", 2)
    add_table(doc, source_tables[4], [1.9, 4.6])

    add_heading(doc, "5. 정형 수학 추출 API", 1)
    p = doc.add_paragraph(style="EAH Note")
    r = p.add_run("화면의 정형 추출기(math-extractor-2)가 사용하는 API와 응답 예시입니다.")
    style_run(r, size=10.2, color=INK)
    add_heading(doc, "API 엔드포인트", 2)
    add_note_box(
        doc,
        "권한 및 구현 위치",
        '비동기 Request-Reply 패턴이며 @PreAuthorize("hasRole(\'ADMIN\')")로 보호됩니다.\n'
        "프론트 클라이언트: structuredMathSheetApi.ts\n"
        "백엔드 컨트롤러: StructuredMathSheetController.java",
    )
    add_table(doc, source_tables[5], [1.3, 3.05, 2.15])
    add_api_flow(doc)

    add_heading(doc, "핵심 타입", 2)
    add_note_box(doc, "타입 매핑", "프론트 StructuredMathItem과 백엔드 StructuredMathSheetResponse.Item은 1:1로 매핑됩니다.")
    add_code_block(
        doc,
        """type StructuredMathItem = {
  questionNumber: number | null;
  prompt: string;                  // 발문 (한글 + $LaTeX$ 인라인)
  choices: string[];               // 보기 ①~⑤, 단답형이면 []
  figureImageUrl: string | null;   // 도형 이미지
  answer: string | null;
  points: number | null;           // 배점 → "[2점]"
  type: string | null;
  subject: string | null;
  needsReview: boolean;
};

public record StructuredMathJobResponse(
  String status,   // "PROCESSING" | "DONE" | "FAILED"
  String sheetId,  // DONE 전엔 null
  String error,    // FAILED 아니면 null
  int total,
  int done
) {}""",
    )

    add_heading(doc, "6. 핵심 도메인 API", 1)
    add_table(doc, source_tables[6], [6.5], header=False)
    add_table(doc, source_tables[7], [1.15, 1.95, 3.4])
    add_heading(doc, "시험 흐름", 2)
    add_note_box(
        doc,
        "문항 등록/추출 → 임베딩/유사 검색 → 시험지 구성(DRAFT) → 게시(PUBLISHED) → 응시 시작/제출 → 자동 채점 → 결과 분석/오답노트 → 마감(CLOSED)",
        "",
        fill="F7FAFF",
    )

    add_heading(doc, "7. 기타 API", 1)
    add_table(doc, source_tables[8], [1.65, 4.85])

    add_heading(doc, "8. 배포/외부 연동", 1)
    add_table(doc, source_tables[9], [1.45, 5.05])

    # Remove the initial blank paragraph from the new document template if present.
    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)
